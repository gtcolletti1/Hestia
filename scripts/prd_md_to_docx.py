#!/usr/bin/env python3
"""Convert the Hestia PRD markdown to a styled .docx.

Usage:
    python3 scripts/prd_md_to_docx.py [SRC_MD] [DST_DOCX]

Defaults to PRD_Family_Hub_Display_v2.md → PRD_Family_Hub_Display_v2.docx
in the repository root. Requires `python-docx` (pip install python-docx).
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt

REPO_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_SRC = REPO_ROOT / "PRD_Family_Hub_Display_v2.md"
DEFAULT_DST = REPO_ROOT / "PRD_Family_Hub_Display_v2.docx"

SRC = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_SRC
DST = Path(sys.argv[2]) if len(sys.argv) > 2 else DEFAULT_DST

INLINE_RE = re.compile(
    r"(\*\*([^*]+)\*\*)"      # bold
    r"|(\*([^*]+)\*)"          # italic
    r"|(`([^`]+)`)"            # code
)


def add_inline_runs(paragraph, text):
    """Parse inline **bold**, *italic*, `code` and append runs."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(1):
            r = paragraph.add_run(m.group(2)); r.bold = True
        elif m.group(3):
            r = paragraph.add_run(m.group(4)); r.italic = True
        elif m.group(5):
            r = paragraph.add_run(m.group(6))
            r.font.name = "Courier New"
            r.font.size = Pt(10)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table(lines, start):
    """Parse a markdown pipe table starting at lines[start]. Returns (rows, end_index_exclusive)."""
    rows = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        rows.append(lines[i])
        i += 1
    # rows[0] = header, rows[1] = separator, rows[2:] = data
    def split_row(row):
        cells = row.strip().strip("|").split("|")
        return [c.strip() for c in cells]
    header = split_row(rows[0])
    data = [split_row(r) for r in rows[2:]] if len(rows) > 2 else []
    return header, data, i


def main():
    md = SRC.read_text().splitlines()
    doc = Document()

    # Default font tweaks
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title_set = False
    i = 0
    while i < len(md):
        line = md[i]
        stripped = line.rstrip()

        # Skip blank lines (paragraph breaks happen naturally)
        if not stripped.strip():
            i += 1
            continue

        # Horizontal rule
        if stripped.strip() == "---":
            i += 1
            continue

        # Code fence
        if stripped.lstrip().startswith("```"):
            lang = stripped.lstrip()[3:].strip()
            i += 1
            code_lines = []
            while i < len(md) and not md[i].lstrip().startswith("```"):
                code_lines.append(md[i])
                i += 1
            i += 1  # skip closing fence
            for cl in code_lines:
                p = doc.add_paragraph()
                r = p.add_run(cl)
                r.font.name = "Courier New"
                r.font.size = Pt(10)
            continue

        # Headings
        if stripped.startswith("# "):
            text = stripped[2:].strip()
            if not title_set:
                # Split "Hestia — Product Requirements Document v2" into Title + H1
                if "—" in text:
                    name, rest = text.split("—", 1)
                    p = doc.add_paragraph(name.strip(), style="Title")
                    doc.add_paragraph(rest.strip(), style="Heading 1")
                else:
                    doc.add_paragraph(text, style="Title")
                title_set = True
            else:
                doc.add_paragraph(text, style="Heading 1")
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_paragraph(stripped[3:].strip(), style="Heading 1")
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_paragraph(stripped[4:].strip(), style="Heading 2")
            i += 1
            continue
        if stripped.startswith("#### "):
            doc.add_paragraph(stripped[5:].strip(), style="Heading 3")
            i += 1
            continue

        # Blockquote (collect consecutive)
        if stripped.lstrip().startswith(">"):
            quote_lines = []
            while i < len(md) and md[i].lstrip().startswith(">"):
                quote_lines.append(md[i].lstrip()[1:].lstrip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            for j, ql in enumerate(quote_lines):
                if j > 0:
                    p.add_run("\n")
                start_idx = len(p.runs)
                add_inline_runs(p, ql)
                for run in p.runs[start_idx:]:
                    run.italic = True
            continue

        # Table
        if stripped.lstrip().startswith("|"):
            header, data, end = parse_table(md, i)
            table = doc.add_table(rows=1 + len(data), cols=len(header))
            try:
                table.style = "Light Grid Accent 1"
            except KeyError:
                table.style = "Table Grid"
            # header
            for c, h in enumerate(header):
                cell = table.rows[0].cells[c]
                cell.text = ""
                add_inline_runs(cell.paragraphs[0], h)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            # data
            for r_idx, row in enumerate(data, start=1):
                for c_idx, val in enumerate(row):
                    if c_idx >= len(table.rows[r_idx].cells):
                        continue
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.text = ""
                    add_inline_runs(cell.paragraphs[0], val)
            i = end
            continue

        # Bullet list (handles "- ", "* ", "- [x] ", "- [ ] ")
        bullet_match = re.match(r"^(\s*)([-*])\s+(.*)$", line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = bullet_match.group(3)
            # Checkbox prefix
            checkbox = ""
            cb = re.match(r"^\[([ xX])\]\s+(.*)$", text)
            if cb:
                checkbox = "☑ " if cb.group(1).lower() == "x" else "☐ "
                text = cb.group(2)
            style_name = "List Bullet" if indent == 0 else "List Bullet 2"
            try:
                p = doc.add_paragraph(style=style_name)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
            if checkbox:
                p.add_run(checkbox)
            add_inline_runs(p, text)
            i += 1
            continue

        # Ordered list
        ol_match = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if ol_match:
            text = ol_match.group(2)
            try:
                p = doc.add_paragraph(style="List Number")
            except KeyError:
                p = doc.add_paragraph()
            add_inline_runs(p, text)
            i += 1
            continue

        # Paragraph (collect consecutive non-blank, non-special lines)
        para_lines = [stripped]
        i += 1
        while i < len(md):
            nxt = md[i].rstrip()
            if not nxt.strip():
                break
            if re.match(r"^(#{1,6} |\||> |\s*[-*]\s|\s*\d+\.\s|```|---$)", nxt):
                break
            para_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_inline_runs(p, " ".join(para_lines))

    doc.save(DST)
    print(f"Wrote {DST} ({DST.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
